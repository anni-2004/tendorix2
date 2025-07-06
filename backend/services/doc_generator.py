import os
import re
from docx import Document
from pymongo import MongoClient
from backend.services.field_mapper import map_fields_by_embedding
from backend.services.schema import load_template_schema

# MongoDB Setup
MONGO_URI = "mongodb+srv://anirudh:anirudh@cluster0.vdtqwjl.mongodb.net/?retryWrites=true&w=majority"
DB_NAME = "tender_system"
COLLECTION_NAME = "tenders"
client = MongoClient(MONGO_URI)
db = client[DB_NAME]
collection = db[COLLECTION_NAME]

def generate_docx_from_template(template_string: str, mapped_data: dict, output_path: str = "output/Final_Generated_Document.docx"):
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    # Normalize curly quotes and braces
    template_string = re.sub(r'[“”]', '"', template_string)
    template_string = re.sub(r"[‘’]", "'", template_string)

    print("📄 Original Template:")
    print(template_string)
    print("🔑 Mapped Data:")
    for k, v in mapped_data.items():
        print(f"  {k}: {v}")

    # Replace placeholders with values
    filled_text = template_string
    for key, value in mapped_data.items():
        placeholder = f"{{{key}}}"
        filled_text = filled_text.replace(placeholder, str(value) if value else "")

    # Clean unreplaced placeholders
    unreplaced = re.findall(r"\{([^}]+)\}", filled_text)
    if unreplaced:
        print("⚠️ Unreplaced placeholders found:", unreplaced)
    filled_text = re.sub(r"\{[^}]+\}", "", filled_text)

    # Write to DOCX
    doc = Document()
    doc.styles['Normal'].paragraph_format.space_after = 0

    for line in filled_text.split("\n"):
        if not line.strip():
            continue
        p = doc.add_paragraph()
        if line.strip().isupper():
            p.add_run(line.strip()).bold = True
        else:
            p.add_run(line.strip())

    doc.save(output_path)
    print(f"\n✅ Document generated at: {output_path}")


# ------------------------
# Example usage
# ------------------------

if __name__ == "__main__":
    # Load schema dynamically from template DOCX
    schema = load_template_schema("Refundable Deposit Refund Request Letters- EMD.docx")
    if not schema:
        raise RuntimeError("❌ Failed to load template schema.")

    template_string = schema["templateString"]
    gemini_fields = schema["fields"]

    # Fetch tender data from MongoDB
    tender_id = "LEGACY11784"  # change dynamically if needed
    backend_data = collection.find_one({"Tender ID": tender_id})

    if not backend_data:
        raise ValueError(f"❌ Tender not found: {tender_id}")

    backend_fields = list(backend_data.keys())

    # Map fields using embedding-based similarity
    mapped_data = map_fields_by_embedding(gemini_fields, backend_fields, backend_data)

    # Generate final document
    generate_docx_from_template(template_string, mapped_data, "output/AZPIL_Final_Document.docx")
